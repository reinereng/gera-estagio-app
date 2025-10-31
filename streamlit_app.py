import streamlit as st
from docx import Document
from docx.shared import Pt

import tempfile
from datetime import datetime, timedelta
import numpy as np
import pandas as pd

from PIL import Image
import requests
from io import BytesIO

# URL da imagem no GitHub
image_url = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/imagens/logo_rede_unirb.jpg"

# Carrega a imagem diretamente da URL
response = requests.get(image_url)
img = Image.open(BytesIO(response.content))

# Caminho das DOCUMENTAÇÕES E IMAGENS

# Função para preencher o documento de termo de compromisso
def preencher_termo(nome_aluno, documento):
    dados_alunos = str(nome_aluno) + ", RG: " + str(rg_aluno) + ", CPF: " + str(cpf_aluno)
    dados_alunos = dados_alunos + ", estudante com matrícula " + str(matricula) + ", residente " + endereco_aluno
    
    dados_empresa = nome_empresa + " com sede e foro na Cidade de " + cidade_empresa + ", " + uf_empresa + ", estabelecida no endereço: "
    dados_empresa = dados_empresa +  endereco_empresa + ", " + str(bairro_empresa) + ", CEP: " + str(cep_empresa) + ", cadastrada no CNPJ: " 
    dados_empresa = dados_empresa + cnpj_empresa + ", neste ato representada por " + representante + ", CPF: " + cpf_empresa
 
    horasestagio = str(total_horas_estagio) + " horas, com jornada de estágio de " + str(horas_por_dia) + " horas por dia e carga horária máxima de 30 horas/semanais"
    
    supervisor_texto = supervisor + ", " + cargosupervisor + ", com registro no conselho profissional sob o número "  + str(conselho_empresa)

    # Usar o documento já carregado
    doc = documento
    
    # Substitui as informações do aluno
    for paragrafo in doc.paragraphs:
        for run in paragrafo.runs:  # Cada 'run' é uma sequência de caracteres com formatação específica  
            if '<DADOS_EMPRESA>' in run.text:
                run.text = run.text.replace('<DADOS_EMPRESA>', dados_empresa)   
                run.bold = True
            if '<DADOS_ALUNO>' in run.text:
                run.text = run.text.replace('<DADOS_ALUNO>', dados_alunos)   
                run.bold = True
            if '<CURSO_ALUNO>' in run.text:
                run.text = run.text.replace('<CURSO_ALUNO>', curso_aluno.upper()) 
                run.bold = True         
            if '<SEMESTRE>' in run.text:
                run.text = run.text.replace('<SEMESTRE>', str(semestre))   
                run.bold = True           
            if '<DATAIN>' in run.text:
                run.text = run.text.replace('<DATAIN>', data_inicio_str)   
                run.bold = True                    
            if '<DATAFIM>' in run.text:
                run.text = run.text.replace('<DATAFIM>', data_termino_str)   
                run.bold = True         
            if '<HORAS>' in run.text:
                run.text = run.text.replace('<HORAS>', horasestagio)   
                run.bold = True                                
            if '<PROFESSOR>' in run.text:
                run.text = run.text.replace('<PROFESSOR>', professor)   
                run.bold = True                                
            if '<SUPERVISOR>' in run.text:
                run.text = run.text.replace('<SUPERVISOR>', supervisor_texto)   
                run.bold = True                                
            if '<ALUNO>' in run.text:
                run.text = run.text.replace('<ALUNO>', nome_aluno)   
                run.bold = True                                
            if '<EMPRESA>' in run.text:
                run.text = run.text.replace('<EMPRESA>', nome_empresa.upper())   
                run.bold = True       
            if '<DONOEMPRESA>' in run.text:
                run.text = run.text.replace('<DONOEMPRESA>', representante)   
                run.bold = True                   
            if '<HOJE>' in run.text:
                run.text = run.text.replace('<HOJE>', datetime.now().strftime('%d/%m/%Y'))   
                run.bold = True      
                   
    # Salva o documento temporariamente
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name


# ----------------------- STREAMLIT ------------------------------------------
# ----------------------------------------------------------------------------
# Exibe a logo no topo da aplicação
st.image(img, width=300)  # Ajuste o valor de width conforme necessário

# Interface do Streamlit
st.title('Gerador de documentação - Rede Unirb')

# Lista de opções de IES
opcoes_ies = [
    "Centro Universitário Unirb",
    "Centro Universitário Unirb Alagoinhas",
    "Faculdade Castro Alves",
    "Faculdade Diplomata",
    "Faculdade Unirb - Arapiraca",
    "Faculdade Unirb - Aracaju",
    "Faculdade Unirb - Barreiras",
    "Faculdade Unirb - Feira de Santana",
    "Faculdade Unirb - Maceió",
    "Faculdade Unirb Parnaíba",
    "Faculdade Unirb Piauí",
    "Faculdade Unirb Teresina"
]

# Exibe a lista suspensa para escolha da IES
ies_escolhida = st.selectbox("Escolha a IES", opcoes_ies)

# Exibe a opção escolhida
st.write(f"Você selecionou: {ies_escolhida}")

# Pergunta ao usuário o tipo de documento a ser gerado
tipo_documento = st.radio(
    "Selecione o tipo de documento:",
    ("Estágio", "Trabalho de Conclução de Curso")
)

# Escolha do método de entrada
opcao_entrada = st.radio(
    "Como você gostaria de inserir os dados?",
    ("Manual", "Planilha")
)
# Lista de opções de IES
opcoes_Professor = [
    " ",
    "Outro",
    "Adriana Barros Dias",
    "Aline Alves Bandeira",
    "Alfredo Jorge Gomes Silva",
    "Anderson Ravanny de Andrade Gomes",
    "Andréa Luciane de Paula Lacerda",
    "Christiane Baracho Pereira da Silva",
    "Cícera Emanuelle Gomes de Oliveira",
    "Daniela Santos Ribeiro",
    "Darlan Francisco Rocha dos Santos",
    "Dilcinéa dos Santos Reis",
    "Dorgilan Rodrigues da Cruz",
    "Eliane Costa dos Santos Baptista",
    "Emilia Maria Modesto de Menezes",
    "Eudilena Laurindo de Medeiros",
    "Érika Souza Vieira",
    "Everton Moraes Lopes",
    "Ivy Góis da Fonsêca Lyra",
    "Jacyara Silva Oliveira",
    "Jandira Dantas dos Santos",
    "José Marques de Vasconcelos Filho",
    "Liliane Machado Nascimento Pereira",
    "Livia Maria Sales Lima Magalhães",
    "Lucas Kayzan Barbosa da Silva",
    "Luiz Sampaio Athayde Junior",
    "Luciana Labidel dos Santos",
    "Marcel Engrácio Leal da Silva",
    "Michelle Diana Leal Pinheiro Matos",
    "Midiã Oliveira Lima",
    "Patrícia Mara Medeiros",
    "Paulo Fernando Araujo Feitosa Leite",
    "Paulo Rogério Menezes de Almeida",
    "Pedro Florencio Ribeiro",
    "Reiner Requião de Souza",
    "Taiana Gomes Libório",
    "Tássia Rangel Guerreiro dos Santos"
]
opcoes_titulo = [
    "Bach.",
    "Esp.",
    "Ma.",
    "Me.",
    "Dra.",
    "Dr."
    ]

# Formatar a data no formato por extenso
meses = {
    1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
    5: "maio", 6: "junho", 7: "julho", 8: "agosto",
    9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
}

# =======================================   TCC   ==========================================
if tipo_documento == "Trabalho de Conclução de Curso":
    
    #DOCUMENTOS
    if ies_escolhida == "Centro Universitário Unirb":
        doc_url_ata  = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_ATA_TCC_Salvador.docx"
        doc_url_ori  = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Declaração_Orientador_Salvador.docx"
        doc_url_resp = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Responsabilidade_Salvador.docx"
        doc_url_auto = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Autorização_Salvador.docx"
    
    elif ies_escolhida == "Centro Universitário Unirb Alagoinhas":
        doc_url_ata = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_ATA_TCC_Alagoinhas.docx"
        doc_url_ori  = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Declaração_Orientador_Alagoinhas.docx"
        doc_url_resp = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Responsabilidade_Alagoinhas.docx"
        doc_url_auto = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Autorização_Alagoinhas.docx"

    elif ies_escolhida == "Faculdade Unirb - Arapiraca":
        doc_url_ata = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_ATA_TCC_Arapiraca.docx"
        doc_url_ori  = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Declaração_Orientador_Arapiraca.docx"
        doc_url_resp = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Responsabilidade_Arapiraca.docx"
        doc_url_auto = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Autorização_Arapiraca.docx"

    elif ies_escolhida == "Faculdade Unirb - Aracaju":
        doc_url_ata = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_ATA_TCC_Aracaju.docx"
        doc_url_ori  = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Declaração_Orientador_Aracaju.docx"
        doc_url_resp = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Responsabilidade_Aracaju.docx"
        doc_url_auto = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Autorização_Aracaju.docx"

    # elif ies_escolhida == "Faculdade Unirb - Feira de Santana":
    #     doc_url_ata = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_ATA_TCC_Feira_de_Santana.docx"
    #     doc_url_ori  = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Declaração_Orientador_Feira_de_Santana.docx"
    #     doc_url_resp = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Responsabilidade_Feira_de_Santana.docx"
    #     doc_url_auto = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Autorização_Feira_de_Santana
    
    elif ies_escolhida == "Faculdade Unirb - Barreiras":
        doc_url_ata = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_ATA_TCC_Barreiras.docx"
        doc_url_ori  = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Declaração_Orientador_Barreiras.docx"
        doc_url_resp = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Responsabilidade_Barreiras.docx"
        doc_url_auto = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Autorização_Barreiras.docx"

    elif ies_escolhida == "Faculdade Unirb - Maceió":
        doc_url_ata = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_ATA_TCC_Maceió.docx"
        doc_url_ori  = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Declaração_Orientador_Maceió.docx"
        doc_url_resp = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Responsabilidade_Maceió.docx"
        doc_url_auto = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Autorização_Maceió.docx"

    elif ies_escolhida == "Faculdade Unirb Parnaíba":
        doc_url_ata = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_ATA_TCC_Parnaíba.docx"
        doc_url_ori  = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Declaração_Orientador_Parnaíba.docx"
        doc_url_resp = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Responsabilidade_Parnaíba.docx"
        doc_url_auto = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Autorização_Parnaíba.docx"

    elif ies_escolhida == "Faculdade Unirb Piauí":
        doc_url_ata = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_ATA_TCC_Piaui.docx"
        doc_url_ori  = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Declaração_Orientador_Piaui.docx"
        doc_url_resp = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Responsabilidade_Piaui.docx"
        doc_url_auto = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Autorização_Piaui.docx"

    elif ies_escolhida == "Faculdade Unirb Teresina":
        doc_url_ata = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_ATA_TCC_Teresina.docx"
        doc_url_ori  = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Declaração_Orientador_Teresina.docx"
        doc_url_resp = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Responsabilidade_Teresina.docx"
        doc_url_auto = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Autorização_Teresina.docx"

    elif ies_escolhida == "Faculdade Diplomata":
        doc_url_ata = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_ATA_TCC_Diplomata.docx"
        doc_url_ori  = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Declaração_Orientador_Diplomata.docx"
        doc_url_resp = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Responsabilidade_Diplomata.docx"
        doc_url_auto = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Autorização_Diplomata.docx"

    elif ies_escolhida == "Faculdade Castro Alves":
        doc_url_ata = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_ATA_TCC_Castro_Alves.docx"
        doc_url_ori  = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Declaração_Orientador_Castro_Alves.docx"
        doc_url_resp = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Responsabilidade_Castro_Alves.docx"
        doc_url_auto = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Termo_de_Autorização_Castro_Alves.docx"

    # Exibe a opção escolhida
    st.write(f"Você selecionou: {ies_escolhida}")
    
    # Inicializa o estado da sessão para armazenar os arquivos gerados e o DataFrame
    if 'arquivos_gerados' not in st.session_state:
        st.session_state.arquivos_gerados = {}

    if 'df' not in st.session_state:
        st.session_state.df = None

    if opcao_entrada == "Planilha":
        # URL do arquivo de exemplo no GitHub
        planilha_url = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Dados_Estagio_Exemplo.xlsx"

        # Faz o download direto do arquivo da URL
        response = requests.get(planilha_url)

        # Botão de download para a planilha de exemplo
        st.subheader("Baixar planilha de exemplo")
        st.write("Clique no botão abaixo para baixar a planilha de exemplo para preenchimento:")
        
    else:
        st.subheader("Dados do Aluno")
        # Cria três colunas para os campos do formulário
        c1, c2 = st.columns(2)
        with c1:
            nome_aluno = st.text_input("Nome do Aluno:")
        with c2:
            matricula = st.text_input("Matrícula:")    
              
        curso_aluno = st.text_input("Curso do Aluno:")

        st.subheader("Dados da Banca")
        # Cria três colunas para os campos do formulário
        titulo = st.text_input("Título do TCC:")
        c3, c4 = st.columns(2)
        with c3:
            data_defesa_aux = st.date_input("Data da defesa da Banca:")
            # Formatar as datas para o formato "DD/MM/YYYY"
            data_defesa = data_defesa_aux.strftime('%d/%m/%Y')
            # Extrair partes da data
            dia = data_defesa_aux.day
            mes = meses[data_defesa_aux.month]
            ano = data_defesa_aux.year

            # Criar a data por extenso
            data_defesa_ext = f"{dia} de {mes} de {ano}"

            Modalidade = st.selectbox("Modalidade", ["GoogleMeet", "Presencial"])
            semestre = st.text_input("Semestre de Orientação (ex.: 2024.2)")
            orientador = st.selectbox("Professor Orientador:", opcoes_Professor)
            banca1 = st.selectbox("Professor 01 da Banca", opcoes_Professor )
            if banca1 == "Outro":
                banca1 = st.text_input("Digite o nome do Professor 01")
            banca2 = st.selectbox("Professor 02 da Banca", opcoes_Professor)
            if banca2 == "Outro":
                banca2 = st.text_input("Digite o nome do Professor 02")

        with c4:
            hora_defesa = st.text_input("Horário da Defesa da Banca:")
            notaTCC = st.text_input("Nota do TCC:")
            st.markdown("<div style='margin-bottom: 85px;'></div>", unsafe_allow_html=True)

            
            formacao00 = st.selectbox("Titulação do Orientador:", opcoes_titulo)
            formacao01 = st.selectbox("Titulação do Professor 01:", opcoes_titulo)
            formacao02 = st.selectbox("Titulação do Professor 02:", opcoes_titulo)

        if 'arquivos_temp' not in st.session_state:
            st.session_state.arquivos_temp = {}

        if st.button("Gerar Todos os Documentos"):
            # Baixar o modelo de documento da web
            response_ata = requests.get(doc_url_ata)
            doc_ata = Document(BytesIO(response_ata.content))
            
            # Substituir marcadores de texto
            texto1 = "Na data de " + data_defesa + ", no horário das " + hora_defesa 
            if Modalidade == "GoogleMeet":
                texto1 = texto1 + ", em reunião virtual via GoogleMeet, "
            elif Modalidade == "Presencial":
                texto1 = texto1 + ", na sede da IES, "
            texto1 = texto1 + "realizou-se a defesa pública do Trabalho de Conclusão de Curso – TCC do discente "             
            texto1 = texto1 + nome_aluno + ", " + matricula + ", intitulado: " + titulo + "."

            paragrafo1 = texto1

            texto4 = "A Banca Examinadora, composta pelos professores " + formacao00 + " " + orientador + " (como presidente e orientador), "
            texto5 = formacao01 + " " + banca1 + " e " + formacao02 + " " + banca2
            texto6 = ", após avaliação e deliberação, considerou o trabalho:"       

            paragrafo2 = texto4 + texto5 + texto6
            
            indices_paragrafos = [1, 3, 5, 9, 16, 20, 23]

            # Iterar sobre os parágrafos
            for j, paragrafos in enumerate(doc_ata.paragraphs):
                if j in indices_paragrafos:                
                    paragrafos.text = paragrafos.text.replace("<<CURSO>>", curso_aluno)
                    paragrafos.text = paragrafos.text.replace("<<paragrafo1>>", paragrafo1)
                    paragrafos.text = paragrafos.text.replace("<<paragrafo2>>", paragrafo2)
                    paragrafos.text = paragrafos.text.replace("<<nota>>", notaTCC)
                    paragrafos.text = paragrafos.text.replace("<<orientador>>", formacao00 + " " + orientador)
                    paragrafos.text = paragrafos.text.replace("<<banca1>>", formacao01 + ". " + banca1)
                    paragrafos.text = paragrafos.text.replace("<<banca2>>", formacao02 + ". " + banca2)
                    
            # Salvando o documento preenchido
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            doc_ata.save(temp_file.name)
            
            response_ori = requests.get(doc_url_ori)
            doc_ori = Document(BytesIO(response_ori.content))
            
            indices_paragrafos = [2, 5, 10]

            # Iterar sobre os parágrafos
            for j, paragrafos in enumerate(doc_ori.paragraphs):
                if j in indices_paragrafos:                
                    paragrafos.text = paragrafos.text.replace("<<aluno>>", nome_aluno)
                    paragrafos.text = paragrafos.text.replace("<<titulo>>", titulo)
                    paragrafos.text = paragrafos.text.replace("<<data>>", data_defesa_ext)
                    paragrafos.text = paragrafos.text.replace("<<orientador>>", formacao00 + " " + orientador)
                    paragrafos.text = paragrafos.text.replace("<<CURSO>>", curso_aluno)
            
            response_resp = requests.get(doc_url_resp)
            doc_resp = Document(BytesIO(response_resp.content))

            indices_paragrafos = [2, 8, 10]

            # Iterar sobre os parágrafos
            for j, paragrafos in enumerate(doc_resp.paragraphs):
                if j in indices_paragrafos:                
                    paragrafos.text = paragrafos.text.replace("<<aluno>>", nome_aluno)
                    paragrafos.text = paragrafos.text.replace("<<titulo>>", titulo)
                    paragrafos.text = paragrafos.text.replace("<<data>>", data_defesa_ext)
                    paragrafos.text = paragrafos.text.replace("<<orientador>>", formacao00 + " " + orientador)
                    paragrafos.text = paragrafos.text.replace("<<CURSO>>", curso_aluno)
                                
            response_auto = requests.get(doc_url_auto)
            doc_auto = Document(BytesIO(response_auto.content))

            indices_paragrafos = [2, 4, 8, 11]

            # Iterar sobre os parágrafos
            for j, paragrafos in enumerate(doc_auto.paragraphs):
                if j in indices_paragrafos:                
                    paragrafos.text = paragrafos.text.replace("<<aluno>>", nome_aluno)
                    paragrafos.text = paragrafos.text.replace("<<titulo>>", titulo)
                    paragrafos.text = paragrafos.text.replace("<<data>>", data_defesa_ext)
                    paragrafos.text = paragrafos.text.replace("<<orientador>>", formacao00 + " " + orientador)
                    paragrafos.text = paragrafos.text.replace("<<CURSO>>", curso_aluno)


                print(f"Parágrafo {j + 1}: {paragrafos.text}")

                    # Salvar documentos temporariamente
               
            for nome_doc, documento in [("Ata", doc_ata), 
                                    ("Declaração_Orientador", doc_ori), 
                                    ("Termo_Responsabilidade", doc_resp), 
                                    ("Termo_Autorização", doc_auto)]:
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                documento.save(temp_file.name)
                st.session_state.arquivos_temp[nome_doc] = temp_file.name
            
        # Mostrar botões de download para cada documento gerado
        if st.session_state.arquivos_temp:
            st.subheader("Faça o download dos documentos:")
            for nome_doc, caminho in st.session_state.arquivos_temp.items():
                with open(caminho, "rb") as file:
                    st.download_button(
                        label=f"Download {nome_doc}",
                        data=file,
                        file_name=f"{nome_doc}_{nome_aluno.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
    # Inicializa a lista de arquivos no estado da sessão, se ainda não existir
    # if 'arquivos_certificados' not in st.session_state:
        # st.session_state.arquivos_certificados = []

    st.subheader("Geração de Certificados")   
    
    doc_url_certOri  = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Certificado_Orientador.docx"
    response_resp = requests.get(doc_url_certOri)
    doc_certOri = Document(BytesIO(response_resp.content))
    
    # Caminho do modelo de certificado
    c10, c20 = st.columns(2)
    with c10:
        gestor1 = st.selectbox("Quem vai assinar do Orientador?", opcoes_Professor)
    with c20:
        gestor2 = st.selectbox("Quem vai assinar dos Professores da Banca", opcoes_Professor)   
    
    if Modalidade == "GoogleMeet":
        textoOr = "em reunião virtual via GoogleMeet."
    elif Modalidade == "Presencial":
        textoOr = "presencialmente na sede da IES. "  
    
    if st.button("Gerar Certificados"):
        st.session_state.arquivos_certificados = []  
        arquivos_certificados = []
        indices_paragrafos = [3, 4, 7, 8, 9]
            
        for j, paragrafos in enumerate(doc_certOri.paragraphs):
            if j in indices_paragrafos:

                # Substituir marcadores de texto
                paragrafos.text = paragrafos.text.replace("<<orientador>>", formacao00 + " " + orientador)
                paragrafos.text = paragrafos.text.replace("<<aluno>>", nome_aluno)
                paragrafos.text = paragrafos.text.replace("<<titulo>>", titulo)
                paragrafos.text = paragrafos.text.replace("<<semestre>>", semestre)
                
                paragrafos.text = paragrafos.text.replace("<<data>>", data_defesa_ext)
                paragrafos.text = paragrafos.text.replace("<<tipo>>", textoOr)
                paragrafos.text = paragrafos.text.replace("<<gestor2>>", gestor1)
                # Iterar sobre as corridas dentro do parágrafo
                for run in paragrafos.runs:
                    # Alterar o tamanho da fonte com tratamento para erros
                    try:
                        run.font.size = Pt(18)  # Substitua 18 pelo tamanho desejado
                        run.font.name = "Arial"  # Substitua "Arial" pela fonte desejada
                        run.bold = True  # Aplicar negrito
                    except AttributeError as e:
                        st.warning(f"Erro ao ajustar a formatação do texto: {e}")
        
        # Construir o nome do arquivo
        nome_arquivo = "Certificado_Orientação_TCC_" + orientador + "_Aluno_"+ nome_aluno + ".docx"
    
        # Salvar o certificado temporariamente
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc_certOri.save(temp_file.name)
        arquivos_certificados.append(temp_file.name)
            
        doc_url_certBanca  = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Python_Certificado_Banca.docx"
        response_resp = requests.get(doc_url_certBanca)
        doc_certB1 = Document(BytesIO(response_resp.content))

        indices_paragrafos = [3, 4, 5, 6, 7, 8, 9]
        
        for j, paragrafos in enumerate(doc_certB1.paragraphs):
            if j in indices_paragrafos:

                # Substituir marcadores de texto
                paragrafos.text = paragrafos.text.replace("<<banca1>>", formacao01 + " " + banca1)
                paragrafos.text = paragrafos.text.replace("<<aluno>>", nome_aluno)
                paragrafos.text = paragrafos.text.replace("<<titulo>>", titulo)

                paragrafos.text = paragrafos.text.replace("<<data>>", data_defesa_ext)
                paragrafos.text = paragrafos.text.replace("<<tipo>>", textoOr)
                paragrafos.text = paragrafos.text.replace("<<gestor1>>", gestor2)

                # Iterar sobre as corridas dentro do parágrafo
                for run in paragrafos.runs:
                    # Alterar o tamanho da fonte com tratamento para erros
                    try:
                        run.font.size = Pt(18)  # Substitua 18 pelo tamanho desejado
                        run.font.name = "Arial"  # Substitua "Arial" pela fonte desejada
                        run.bold = True  # Aplicar negrito
                    except AttributeError as e:
                        st.warning(f"Erro ao ajustar a formatação do texto: {e}")
        
        # Construir o nome do arquivo
        nome_arquivo = "Certificado_Banca_TCC_" + banca1 + "_Aluno_"+ nome_aluno + ".docx"
    
        # Salvar o certificado temporariamente
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc_certB1.save(temp_file.name)
        arquivos_certificados.append(temp_file.name)
        
        doc_certB2 = Document(BytesIO(response_resp.content))
        
        for j, paragrafos in enumerate(doc_certB1.paragraphs):
            if j in indices_paragrafos:

                # Substituir marcadores de texto
                paragrafos.text = paragrafos.text.replace("<<banca1>>", formacao02 + " " + banca2)
                paragrafos.text = paragrafos.text.replace("<<aluno>>", nome_aluno)
                paragrafos.text = paragrafos.text.replace("<<titulo>>", titulo)

                paragrafos.text = paragrafos.text.replace("<<data>>", data_defesa_ext)
                paragrafos.text = paragrafos.text.replace("<<tipo>>", textoOr)
                paragrafos.text = paragrafos.text.replace("<<gestor1>>", gestor2)

                # Iterar sobre as corridas dentro do parágrafo
                for run in paragrafos.runs:
                    # Alterar o tamanho da fonte com tratamento para erros
                    try:
                        run.font.size = Pt(18)  # Substitua 18 pelo tamanho desejado
                        run.font.name = "Arial"  # Substitua "Arial" pela fonte desejada
                        run.bold = True  # Aplicar negrito
                    except AttributeError as e:
                        st.warning(f"Erro ao ajustar a formatação do texto: {e}")
        
        # Construir o nome do arquivo
        nome_arquivo = "Certificado_Banca_TCC_" + banca2 + "_Aluno_"+ nome_aluno + ".docx"
    
        # Salvar o certificado temporariamente
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc_certB2.save(temp_file.name)
        arquivos_certificados.append(temp_file.name)
                    
        st.session_state.arquivos_certificados.extend(arquivos_certificados)
        
    # Mostrar botões de download para cada documento gerado
    if 'arquivos_certificados' in st.session_state:
        if st.session_state.arquivos_certificados:  # Verificar se há arquivos na lista
            st.subheader("Faça o download dos documentos:")
            for idx, caminho in enumerate(st.session_state.arquivos_certificados):
                with open(caminho, "rb") as file:
                    st.download_button(
                        label=f"Download Certificado {idx + 1}",
                        data=file,
                        file_name=f"Certificado_{idx + 1}_{nome_aluno.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        else:
            st.warning("Nenhum arquivo foi gerado. Verifique os dados fornecidos.")

# ======================================= ESTÁGIO ==========================================
if tipo_documento == "Estágio":

    #DOCUMENTOS
    if ies_escolhida == "Centro Universitário Unirb":
        # URL do documento no GitHub para o termo de compromisso e convênio
        doc_url_termo = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Termo_Compromisso_Centro.docx"
        doc_url_conv = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Termo_Convenio_Centro.docx"

    elif ies_escolhida == "Centro Universitário Unirb Alagoinhas":
        # URL do documento no GitHub para o termo de compromisso e convênio
        doc_url_termo = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Termo_Compromisso_Centro_Alagoinhas.docx"
        doc_url_conv = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Termo_Convenio_Centro_Alagoinhas.docx"

    elif ies_escolhida == "Faculdade Diplomata":
        # URL do documento no GitHub para o termo de compromisso e convênio
        doc_url_termo = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Termo_Compromisso_Diplomata.docx"
        doc_url_conv = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Termo_Convenio_Diplomata.docx"

    elif ies_escolhida == "Faculdade Castro Alves":
        # URL do documento no GitHub para o termo de compromisso e convênio
        doc_url_termo = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Termo_Compromisso_Castro_Alves.docx"
        doc_url_conv = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Termo_Convenio_Castro_Alves.docx"

    elif ies_escolhida == "Faculdade Unirb - Barreiras":
        # URL do documento no GitHub para o termo de compromisso e convênio
        doc_url_termo = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Termo_Compromisso_Barreiras.docx"
        doc_url_conv = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Termo_Convenio_Barreiras.docx"

    elif ies_escolhida == "Faculdade Unirb Parnaíba":
        # URL do documento no GitHub para o termo de compromisso e convênio
        doc_url_termo = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Termo_Compromisso_Parnaíba.docx"
        doc_url_conv = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Termo_Convenio_Parnaíba.docx"

    elif ies_escolhida == "Faculdade Unirb Piauí":
        # URL do documento no GitHub para o termo de compromisso e convênio
        doc_url_termo = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Termo_Compromisso_Piaui.docx"
        doc_url_conv = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Modelo_Termo_Convenio_Piaui.docx"

    response_termo = requests.get(doc_url_termo)
    caminho_termo = Document(BytesIO(response_termo.content))

    response_conv = requests.get(doc_url_conv)
    caminho_conv = Document(BytesIO(response_conv.content))

    dados_ies = "AMERICA EDUCACIONAL S.A - CENTRO UNIVERSITÁRIO UNIRB - SALVADOR, situado à Av. Tamburugy, 474 - Patamares, Salvador - BA, CEP: 41680-440, tel: (71) 3368-8300 e-mail: unirb@unirb.edu.br inscrita no CNPJ 28.844.791/0001-91 representada neste ato, por Carlos Joel Pereira, CPF: 159.659.615-53"

    # Exibe a lista suspensa para escolha da Professor
    professor_escolhido = st.selectbox("Escolha o Orientador do Estágio", opcoes_Professor)
    
    if professor_escolhido == "Reiner Requião de Souza":
        professor = "Reiner Requião de Souza, portador do CPF n.º 009.893.855-07 e RG nº 07584711-65 SSP/BA"
    if professor_escolhido == "Adriana Barros Dias":
        professor = "Adriana Barros Dias, portador do CPF n.º 682.516.455-49 e RG nº 04.945.966-00 SSP/BA"
    if professor_escolhido == "Eudilena Laurindo de Medeiros":
        professor = "Eudilena Laurindo de Medeiros, portador do CPF n.º 082.488.904-57 e RG nº 00.022.964-78 SSP/RN"
    if professor_escolhido == "Anderson Ravanny de Andrade Gomes":
        professor = "Anderson Ravanny de Andrade Gomes, portador do CPF n.º 074.662.624-06"
    if professor_escolhido == "Pedro Florencio Ribeiro":
        professor = "Pedro Florencio Ribeiro, portador do CPF n.º 032.775.163-00"
    if professor_escolhido == "Patrícia Medeiros":
        professor = "Patrícia Mara Medeiros, portador do CPF n.º 165.469.298-05"

    # Inicializa o estado da sessão para armazenar os arquivos gerados e o DataFrame
    if 'arquivos_gerados' not in st.session_state:
        st.session_state.arquivos_gerados = {}

    if 'df' not in st.session_state:
        st.session_state.df = None

    if opcao_entrada == "Planilha":
        # URL do arquivo de exemplo no GitHub
        planilha_url = "https://raw.githubusercontent.com/reinereng/gera-estagio-app/main/modelos/Dados_Estagio_Exemplo.xlsx"

        # Faz o download direto do arquivo da URL
        response = requests.get(planilha_url)

        # Botão de download para a planilha de exemplo
        st.subheader("Baixar planilha de exemplo")
        st.write("Clique no botão abaixo para baixar a planilha de exemplo para preenchimento:")

        st.download_button(
            label="Baixar Planilha de Exemplo",
            data=response.content,  # Conteúdo direto da resposta
            file_name="Dados_Estagio_Exemplo.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        st.subheader("Upload da Planilha com Dados")
        # Upload de planilha
        arquivo_planilha = st.file_uploader("Carregue a planilha de entrada conforme a planilha Exemplo", type=["xlsx"])

        if arquivo_planilha is not None:
            # Limpa o DataFrame antigo no estado da sessão
            st.session_state.df = pd.read_excel(arquivo_planilha)
            st.write("Dados carregados da planilha:", st.session_state.df)

            # Faz uma cópia local do DataFrame
            df = st.session_state.df.copy()
            
            # Itera sobre cada linha do DataFrame
            for i in range(len(df)):
                # Dados do Aluno
                nome_aluno = df["Nome do Aluno"][i]
                curso_aluno = df["Curso do Aluno"][i]
                rg_aluno = df["RG do Aluno"][i]
                semestre = df["Semestre Atual do Aluno"][i]
                cpf_aluno = df["CPF do Aluno"][i]
                matricula = df["Matrícula"][i]
                endereco_aluno = df["Endereço Completo do Aluno"][i]
                
                # Dados do Estágio
                total_horas_estagio = df["Total de Horas do Estágio"][i]
                horas_por_dia = df["Horas de Estágio por Dia"][i]

                    
                # Converte data_inicio para datetime com tratamento de erros
                data_inicio = pd.to_datetime(df["Data de Início do Estágio"][i], format='%d/%m/%Y', errors='coerce')
                
                # Converte para o tipo datetime64[D] necessário para np.busday_offset
                if pd.notnull(data_inicio):
                    data_inicio_np = data_inicio.to_numpy().astype('datetime64[D]')
                else:
                    print(f"Data de início inválida para o índice {i}.")
                    continue

                # Verifica se a data de término é nula ou precisa ser calculada
                data_termino = df["Data de Término do Estágio"][i]
                if pd.isnull(data_termino) or data_termino == '':
                    dias_necessarios = np.ceil(total_horas_estagio / horas_por_dia)
                    data_termino_np = np.busday_offset(data_inicio_np, dias_necessarios-1, roll='forward')
                else:
                    # Converte data_termino para datetime
                    data_termino = pd.to_datetime(data_termino, format='%d/%m/%Y', errors='coerce')
                    data_termino_np = data_termino.to_numpy().astype('datetime64[D]')

                # Garante que ambas as datas estejam no formato datetime antes de usar strftime
                if pd.notnull(data_inicio) and pd.notnull(data_termino_np):
                    data_inicio_str = data_inicio.strftime('%d/%m/%Y')
                    data_termino_str = pd.Timestamp(data_termino_np).strftime('%d/%m/%Y')

                    # Processa ou exibe os dados como necessário
                    print(f"Data de Início: {data_inicio_str}, Data de Término: {data_termino_str}")
                else:
                    print("Erro: Datas inválidas encontradas.")    
                    
                
                # Dados da Empresa
                nome_empresa = df["Nome da Empresa"][i]
                representante = df["Representante da Empresa"][i]
                supervisor = df["Supervisor"][i]
                cargosupervisor = df["Cargo do Supervisor"][i]
                cnpj_empresa = df["CNPJ da Empresa"][i]
                cpf_empresa = df["CPF do Representante"][i]
                conselho_empresa = df["Registro do Conselho do Supervisor"][i]
                endereco_empresa = df["Endereço da Empresa"][i]
                bairro_empresa = df["Bairro da Empresa"][i]
                cidade_empresa = df["Cidade da Empresa"][i]
                cep_empresa = df["CEP da Empresa"][i]
                uf_empresa = df["UF da Empresa"][i]

                # Inicializa o estado da sessão para armazenar os arquivos gerados
                if nome_aluno not in st.session_state.arquivos_gerados:
                    st.session_state.arquivos_gerados[nome_aluno] = {'termo': None, 'convenio': None}

                if nome_aluno:
                    # Preenche o termo de compromisso
                    nome_arquivo_termo = preencher_termo(nome_aluno, caminho_termo)
                    
                    # Preenche o termo de convênio
                    nome_arquivo_convenio = preencher_termo(nome_aluno, caminho_conv)
                    
                    # Armazena os caminhos dos arquivos no estado da sessão
                    st.session_state.arquivos_gerados[nome_aluno]['termo'] = nome_arquivo_termo
                    st.session_state.arquivos_gerados[nome_aluno]['convenio'] = nome_arquivo_convenio

                # Verifica se os arquivos foram gerados e exibe os botões de download
                if st.session_state.arquivos_gerados[nome_aluno]['termo'] and st.session_state.arquivos_gerados[nome_aluno]['convenio']:
                    with open(st.session_state.arquivos_gerados[nome_aluno]['termo'], "rb") as file_termo:
                        st.download_button(
                            label=f"Download Termo de Compromisso - {nome_aluno}", 
                            data=file_termo, 
                            file_name=f"Termo_Compromisso_{nome_aluno.replace(' ', '_')}.docx"
                        )
                    
                    with open(st.session_state.arquivos_gerados[nome_aluno]['convenio'], "rb") as file_convenio:
                        st.download_button(
                            label=f"Download Termo de Convênio - {nome_aluno}", 
                            data=file_convenio, 
                            file_name=f"Termo_Convenio_{nome_aluno.replace(' ', '_')}.docx"
                        )

    else:
        st.subheader("Dados do Aluno")
        # Cria três colunas para os campos do formulário
        c1, c2, c3 = st.columns(3)
        with c1:
            nome_aluno = st.text_input("Nome do Aluno")
            curso_aluno = st.text_input("Curso do Aluno")
        with c2:
            rg_aluno = st.text_input("RG do Aluno")
            semestre = st.text_input("Semestre Atual do Aluno")    
        with c3:
            cpf_aluno = st.text_input("CPF do Aluno")
            matricula = st.text_input("Matrícula")    

        endereco_aluno = st.text_input("Endereço Completo do Aluno")    

        st.subheader("Dados do Estágio")
        # Cria duas colunas para os campos do formulário
        col1, col2 = st.columns(2)
        with col1:
            total_horas_estagio = st.number_input("Total de Horas do Estágio", min_value=1, value=120)

        with col2:
            horas_por_dia = st.number_input("Horas de Estágio por Dia", min_value=1, max_value=6, value=4)# Calcula a data de término 20 dias úteis após a data de início

        # Cria duas colunas para os campos do formulário
        col1a, col2a = st.columns(2)
        with col1a:
            data_inicio = st.date_input("Data de Início do Estágio")
            # Calcula o número de dias úteis necessários para completar o estágio
            dias_necessarios = np.ceil(total_horas_estagio / horas_por_dia)

            # Exibe o número de dias úteis calculados
            st.write(f"Número de dias úteis necessários: {int(dias_necessarios)}")

            data_termino_sugerida = np.busday_offset(data_inicio, dias_necessarios-1, roll='forward')

            # Converte a data de término sugerida para o formato datetime do Python
            data_termino_sugerida = datetime.strptime(str(data_termino_sugerida), '%Y-%m-%d')

        with col2a:
            data_termino = st.date_input("Data de Término do Estágio", value=data_termino_sugerida)

        # Formatar as datas para o formato "DD/MM/YYYY"
        data_inicio_str = data_inicio.strftime('%d/%m/%Y')
        data_termino_str = data_termino.strftime('%d/%m/%Y')

        # Calcula o número de dias úteis entre as duas datas
        dias_uteis = np.busday_count(data_inicio, data_termino)

        # Exibe o número de dias úteis
        st.write(f"Total de dias úteis alocados: {dias_uteis}")

        st.subheader("Dados da Empresa")    

        # Cria duas colunas para os campos do formulário
        col6, col7 = st.columns(2)
        with col6:
            nome_empresa = st.text_input("Nome da Empresa")
            representante = st.text_input("Representante da Empresa")    
            supervisor = st.text_input("Supervisor")
            cargosupervisor = st.text_input("Cargo do Supervisor")
        with col7:
            cnpj_empresa = st.text_input("CNPJ da Empresa")    
            cpf_empresa = st.text_input("CPF do Representante")
            conselho_empresa = st.text_input("Registro do Conselho do Supervisor") 

        # Cria três colunas para os campos do formulário
        col3, col4, col5 = st.columns(3)

        # Campos de entrada para os dados do aluno
        with col3:
            endereco_empresa = st.text_input("Endereço da Empresa")
            bairro_empresa = st.text_input("Bairro da Empresa")
        with col4:
            cidade_empresa = st.text_input("Cidade da Empresa")
            cep_empresa = st.text_input("CEP da Empresa")
        with col5:
            uf_empresa = st.text_input("UF da Empresa")
            
        # Botão para gerar o documento
        if st.button("Gerar Documento"):
            
            # Inicializa o estado da sessão para armazenar os arquivos gerados
            if nome_aluno not in st.session_state.arquivos_gerados:
                st.session_state.arquivos_gerados[nome_aluno] = {'termo': None, 'convenio': None}

            if nome_aluno:        
                # Inicializa o estado da sessão para armazenar os arquivos gerados
                nome_arquivo_termo = preencher_termo(nome_aluno, caminho_termo)
                
                # Preenche o termo de convênio
                nome_arquivo_convenio = preencher_termo(nome_aluno, caminho_conv)
                
                # Armazena os caminhos dos arquivos no estado da sessão
                st.session_state.arquivos_gerados[nome_aluno]['termo'] = nome_arquivo_termo
                st.session_state.arquivos_gerados[nome_aluno]['convenio'] = nome_arquivo_convenio

                # Verifica se os arquivos foram gerados e exibe os botões de download
                if st.session_state.arquivos_gerados[nome_aluno]['termo'] and st.session_state.arquivos_gerados[nome_aluno]['convenio']:
                    with open(st.session_state.arquivos_gerados[nome_aluno]['termo'], "rb") as file_termo:
                        st.download_button(
                            label=f"Download Termo de Compromisso - {nome_aluno}", 
                            data=file_termo, 
                            file_name=f"Termo_Compromisso_{nome_aluno.replace(' ', '_')}.docx"
                        )
                    
                    with open(st.session_state.arquivos_gerados[nome_aluno]['convenio'], "rb") as file_convenio:
                        st.download_button(
                            label=f"Download Termo de Convênio - {nome_aluno}", 
                            data=file_convenio, 
                            file_name=f"Termo_Convenio_{nome_aluno.replace(' ', '_')}.docx"
                        )        
            else:
                st.warning("Por favor, insira o nome do aluno.")

